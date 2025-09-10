import streamlit as st
import requests
from docx import Document
import re
import tempfile
import os
import json
import pandas as pd
from typing import List, Dict

# Configuration de la page
st.set_page_config(
    page_title="AI Vulnerability Alignment Auditor",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS
st.markdown("""
<style>
    .header { text-align: center; color: #1E88E5; margin-bottom: 2rem; }
    .vuln-card { 
        background-color: #f8f9fa; padding: 15px; border-radius: 10px; 
        border-left: 5px solid #6c757d; margin: 10px 0; 
    }
    .consistent { border-left: 5px solid #28a745; background-color: #E8F5E9; }
    .partial { border-left: 5px solid #ffc107; background-color: #FFF3CD; }
    .inconsistent { border-left: 5px solid #dc3545; background-color: #FFEBEE; }
    .section-box { 
        background-color: white; padding: 12px; border-radius: 6px; 
        border: 1px solid #dee2e6; margin: 8px 0; font-size: 14px;
    }
    .score-high { color: #28a745; font-weight: bold; font-size: 24px; }
    .score-medium { color: #ffc107; font-weight: bold; font-size: 24px; }
    .score-low { color: #dc3545; font-weight: bold; font-size: 24px; }
</style>
""", unsafe_allow_html=True)

def test_ollama_connection():
    """Test Ollama connection"""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=10)
        return response.status_code == 200
    except:
        return False

def extract_all_vulnerabilities(docx_path: str) -> List[Dict]:
    """Extract all vulnerabilities from DOCX"""
    doc = Document(docx_path)
    vulnerabilities = []
    current_vuln = {}
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.search(r'V_\d{2,3}', text):
            if current_vuln:
                vulnerabilities.append(current_vuln)
            current_vuln = {'id': re.search(r'(V_\d{2,3})', text).group(1), 'title': text}
    
    # Process tables for content
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                left_text = row.cells[0].text.strip()
                right_text = row.cells[1].text.strip()
                
                vuln_match = re.search(r'(V_\d{2,3})', left_text) or re.search(r'(V_\d{2,3})', right_text)
                if vuln_match:
                    current_id = vuln_match.group(1)
                    current_vuln = next((v for v in vulnerabilities if v['id'] == current_id), None)
                    if current_vuln:
                        if vuln_match.group(1) in left_text:
                            current_vuln['title'] = right_text
                        else:
                            current_vuln['title'] = left_text
                
                elif current_vuln:
                    sections = {
                        'constats': ['constats', 'findings'],
                        'preuves': ['preuves', 'evidence'],
                        'recommandation': ['recommandation', 'recommendation'],
                        'impacts': ['impacts', 'consequence'],
                        'niveau_criticite': ['niveau de criticité', 'criticité']
                    }
                    for field, keywords in sections.items():
                        if any(kw in left_text.lower() for kw in keywords):
                            current_vuln[field] = right_text
                            break
    
    if current_vuln:
        vulnerabilities.append(current_vuln)
    
    return [v for v in vulnerabilities if any(v.get(k) for k in ['constats', 'preuves', 'recommandation'])]

def analyze_vulnerability_alignment(vuln_data: Dict, model: str = "llama3:latest") -> Dict:
    """AI analysis of vulnerability information alignment"""
    
    prompt = f"""
CRITICAL VULNERABILITY ALIGNMENT ANALYSIS - SECURITY EXPERT MODE

VULNERABILITY: {vuln_data.get('id')} - {vuln_data.get('title')}

DATA TO ANALYZE:
- FINDINGS: {vuln_data.get('constats', 'Not provided')}
- EVIDENCE: {vuln_data.get('preuves', 'Not provided')}
- IMPACTS: {vuln_data.get('impacts', 'Not provided')}
- CRITICALITY: {vuln_data.get('niveau_criticite', 'Not provided')}
- RECOMMENDATION: {vuln_data.get('recommandation', 'Not provided')}

MISSION: Analyze ALIGNMENT and CONSISTENCY between all information elements.

SPECIFIC CHECKS:
1. 🎯 EVIDENCE-FINDINGS ALIGNMENT: Do the evidence properly support and prove the findings?
2. 🔧 RECOMMENDATION-FINDINGS ALIGNMENT: Does the recommendation correctly address the root cause identified in findings?
3. ⚖️ CRITICALITY-IMPACTS ALIGNMENT: Is the criticality level appropriate given the described impacts?
4. 📊 OVERALL CONSISTENCY: Is all information logically consistent and coherent?

RESPOND STRICTLY in JSON format:
{{
  "vulnerability_id": "{vuln_data.get('id')}",
  "alignment_verdict": "FULLY_ALIGNED|PARTIALLY_ALIGNED|MISALIGNED",
  "alignment_score": 0-100,
  "summary": "Brief overall assessment",
  "detailed_analysis": {{
    "evidence_findings_alignment": {{
      "verdict": "ALIGNED|PARTIAL|MISALIGNED",
      "score": 0-100,
      "explanation": "Detailed analysis of evidence supporting findings"
    }},
    "recommendation_findings_alignment": {{
      "verdict": "ALIGNED|PARTIAL|MISALIGNED", 
      "score": 0-100,
      "explanation": "Detailed analysis of recommendation addressing findings"
    }},
    "criticality_impacts_alignment": {{
      "verdict": "ALIGNED|PARTIAL|MISALIGNED",
      "score": 0-100,
      "explanation": "Detailed analysis of criticality matching impacts"
    }}
  }},
  "strengths": ["strength1", "strength2"],
  "improvement_opportunities": ["opportunity1", "opportunity2"],
  "risk_level": "LOW|MEDIUM|HIGH|CRITICAL"
}}

Be rigorous, technical, and constructive. Focus on alignment quality.
"""

    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "format": "json"
            },
            timeout=180
        )
        
        if response.status_code == 200:
            result = response.json()
            response_text = result.get("response", "").strip()
            
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            return {"error": "JSON not found in response"}
        return {"error": f"API error: {response.status_code}"}
        
    except Exception as e:
        return {"error": f"Analysis failed: {str(e)}"}

def display_alignment_results(analysis_results: List[Dict]):
    """Display all alignment analysis results"""
    
    st.markdown("## 📊 Comprehensive Vulnerability Alignment Report")
    
    # Summary statistics
    total_vulns = len(analysis_results)
    aligned_count = sum(1 for r in r.get('alignment_verdict', '') == 'FULLY_ALIGNED' for r in analysis_results if 'error' not in r)
    avg_score = sum(r.get('alignment_score', 0) for r in analysis_results if 'error' not in r) / max(total_vulns, 1)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Vulnerabilities", total_vulns)
    with col2:
        st.metric("Fully Aligned", f"{aligned_count}/{total_vulns}")
    with col3:
        st.metric("Average Alignment Score", f"{avg_score:.1f}/100")
    
    # Detailed results for each vulnerability
    for result in analysis_results:
        if 'error' in result:
            st.error(f"❌ {result['vulnerability_id']} - Analysis failed: {result['error']}")
            continue
            
        verdict = result.get('alignment_verdict', 'UNKNOWN')
        score = result.get('alignment_score', 0)
        
        # Determine styling
        if verdict == 'FULLY_ALIGNED':
            css_class = 'consistent'
            score_class = 'score-high'
        elif verdict == 'PARTIALLY_ALIGNED':
            css_class = 'partial' 
            score_class = 'score-medium'
        else:
            css_class = 'inconsistent'
            score_class = 'score-low'
        
        st.markdown(f"""
        <div class="vuln-card {css_class}">
            <h3>🔍 {result['vulnerability_id']} - Alignment Analysis</h3>
            <p><strong>Verdict:</strong> {verdict} | <strong>Score:</strong> <span class="{score_class}">{score}/100</span></p>
            <p><strong>Summary:</strong> {result.get('summary', '')}</p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander(f"📋 Detailed Analysis - {result['vulnerability_id']}"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 🎯 Evidence-Findings Alignment")
                detail = result['detailed_analysis']['evidence_findings_alignment']
                st.metric("Score", f"{detail['score']}/100")
                st.info(f"**Verdict:** {detail['verdict']}")
                st.write(detail['explanation'])
                
                st.markdown("### 🔧 Recommendation-Findings Alignment")
                detail = result['detailed_analysis']['recommendation_findings_alignment']
                st.metric("Score", f"{detail['score']}/100")
                st.info(f"**Verdict:** {detail['verdict']}")
                st.write(detail['explanation'])
            
            with col2:
                st.markdown("### ⚖️ Criticality-Impacts Alignment")
                detail = result['detailed_analysis']['criticality_impacts_alignment']
                st.metric("Score", f"{detail['score']}/100")
                st.info(f"**Verdict:** {detail['verdict']}")
                st.write(detail['explanation'])
                
                st.markdown("### 📊 Risk Assessment")
                st.warning(f"**Risk Level:** {result.get('risk_level', 'UNKNOWN')}")
                
                st.markdown("### ✅ Strengths")
                for strength in result.get('strengths', []):
                    st.success(f"• {strength}")
                
                st.markdown("### 📝 Improvement Opportunities")
                for opportunity in result.get('improvement_opportunities', []):
                    st.error(f"• {opportunity}")
    
    # Export all results
    if analysis_results:
        result_json = json.dumps(analysis_results, ensure_ascii=False, indent=2)
        st.download_button(
            "💾 Download Full Alignment Report",
            result_json,
            "vulnerability_alignment_report.json",
            "application/json"
        )

# Main Application
st.markdown('<h1 class="header">🔍 AI Vulnerability Alignment Auditor</h1>', unsafe_allow_html=True)
st.markdown("***Automated consistency checking for vulnerability information alignment***")

# Ollama check
if not test_ollama_connection():
    st.error("❌ Ollama not available. Please start Ollama: `ollama serve`")
    st.stop()

st.success("✅ Ollama connected and ready for alignment analysis!")

# File upload
uploaded_file = st.file_uploader("📤 Upload Audit Report DOCX", type="docx")
if not uploaded_file:
    st.info("Please upload a DOCX audit report to begin alignment analysis")
    st.stop()

# Process file
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    tmp_file.write(uploaded_file.getvalue())
    tmp_path = tmp_file.name

try:
    # Extract vulnerabilities
    with st.spinner("🔍 Extracting vulnerabilities from document..."):
        vulnerabilities = extract_all_vulnerabilities(tmp_path)
    
    if not vulnerabilities:
        st.error("No vulnerabilities found in the document")
        st.stop()
    
    st.success(f"✅ Found {len(vulnerabilities)} vulnerabilities for alignment analysis")
    
    # Show vulnerability overview
    st.markdown("## 📋 Vulnerabilities Found")
    vuln_df = pd.DataFrame([{
        'ID': v['id'],
        'Title': v.get('title', '')[:100] + '...',
        'Has Findings': '✅' if v.get('constats') else '❌',
        'Has Evidence': '✅' if v.get('preuves') else '❌', 
        'Has Recommendation': '✅' if v.get('recommandation') else '❌'
    } for v in vulnerabilities])
    
    st.dataframe(vuln_df, use_container_width=True)
    
    # Start alignment analysis
    if st.button("🚀 Analyze All Vulnerabilities Alignment", type="primary"):
        analysis_results = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, vuln in enumerate(vulnerabilities):
            status_text.text(f"Analyzing {vuln['id']} ({i+1}/{len(vulnerabilities)})...")
            progress_bar.progress((i + 1) / len(vulnerabilities))
            
            result = analyze_vulnerability_alignment(vuln)
            result['vulnerability_id'] = vuln['id']
            result['vulnerability_title'] = vuln.get('title', '')
            analysis_results.append(result)
            
            time.sleep(1)  # Rate limiting
        
        progress_bar.empty()
        status_text.empty()
        
        # Display results
        display_alignment_results(analysis_results)

finally:
    os.unlink(tmp_path)