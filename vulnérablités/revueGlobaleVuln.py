import streamlit as st
import requests
from docx import Document
import re
import tempfile
import os
from pathlib import Path

# Configuration de la page
st.set_page_config(
    page_title="Analyseur de Rapport d'Audit",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 2rem;
    }
    .success-box {
        background-color: #E8F5E9;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #4CAF50;
        margin: 20px 0;
    }
    .analysis-section {
        background-color: #F5F5F5;
        padding: 25px;
        border-radius: 10px;
        margin: 15px 0;
        border: 1px solid #E0E0E0;
    }
    .vulnerability-card {
        background-color: #FFF3E0;
        padding: 15px;
        margin: 10px 0;
        border-radius: 8px;
        border-left: 4px solid #FF9800;
    }
    .stButton>button {
        background-color: #1E88E5;
        color: white;
        font-weight: bold;
        border: none;
        padding: 12px 24px;
        border-radius: 8px;
        width: 100%;
    }
    .stButton>button:hover {
        background-color: #1565C0;
    }
</style>
""", unsafe_allow_html=True)

def clean_text(text):
    """Nettoie le texte en supprimant les espaces excessifs."""
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r' *\n *', '\n', text)
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extrait et nettoie tout le texte d'un fichier DOCX."""
    doc = Document(docx_path)
    all_text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            cleaned = clean_text(paragraph.text)
            all_text.append(cleaned)

    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_cell = clean_text(cell.text)
                    row_cells.append(cleaned_cell)
            if row_cells:
                table_lines.append(" | ".join(row_cells))
        if table_lines:
            all_text.append("\n".join(table_lines))

    return "\n".join(all_text)

def analyze_with_ollama(vulnerabilities_text, model_choice):
    """Envoie le texte à Ollama pour analyse."""
    prompt = f"""
    Tu es un expert en cybersécurité. Analyse cette liste de vulnérabilités en FRANÇAIS.

    FOURNIS UN RAPPORT STRUCTURÉ AVEC:
    1. 📊 **Synthèse du risque global** (2-3 phrases)
    2. 🎯 **Top 5 des vulnérabilités critiques** avec justification
    3. 📋 **Catégorisation par type** (Configuration, Patch, Réseau, etc.)
    4. ✅ **Recommendations prioritaires** (actions concrètes)

    FORMAT ATTENDU:
    Utilise des titres ##, des listes à puces • et sois concis.

    DONNÉES À ANALYSER:
    {vulnerabilities_text[:8000]}  # Limite pour éviter les timeouts
    """

    payload = {
        "model": model_choice,
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post("http://localhost:11434/api/generate", json=payload, timeout=None)
        response.raise_for_status()
        return response.json().get("response", "Erreur: Aucune réponse reçue.")
    except Exception as e:
        return f"Erreur lors de l'analyse: {str(e)}"

# Interface principale
st.markdown('<h1 class="main-header">🔍 Analyseur de Vulnérabilités Cybersecurity</h1>', unsafe_allow_html=True)

# Sidebar pour les paramètres
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/security-checked.png", width=100)
    st.title("Configuration")
    
    model_choice = st.selectbox(
        "Modèle Ollama à utiliser:",
        ["ALIENTELLIGENCE/cybersecuritythreatanalysisv2:latest", "llama3:latest"],
        help="Choisissez le modèle spécialisé cybersecurity pour de meilleurs résultats"
    )
    
    st.info("""
    **Instructions:**
    1. Téléchargez votre rapport DOCX
    2. L'analyse commencera automatiquement
    3. Les résultats s'afficheront ci-contre
    """)

# Zone de téléchargement du fichier
uploaded_file = st.file_uploader(
    "📤 Téléchargez votre rapport d'audit (DOCX)",
    type="docx",
    help="Le document doit contenir la section des vulnérabilités"
)

if uploaded_file is not None:
    # Sauvegarde temporaire du fichier
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    try:
        # Extraction du texte
        with st.spinner("🔍 Extraction du texte en cours..."):
            extracted_text = extract_text_from_docx(tmp_path)
        
        # Affichage des métriques
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Caractères extraits", f"{len(extracted_text):,}")
        with col2:
            st.metric("Mots", f"{len(extracted_text.split()):,}")
        with col3:
            st.metric("Lignes", extracted_text.count('\n') + 1)

        # Bouton d'analyse
        if st.button("🚀 Lancer l'Analyse IA", type="primary"):
            with st.spinner("🤖 Analyse en cours par l'IA (2-3 minutes)..."):
                analysis_result = analyze_with_ollama(extracted_text, model_choice)
            
            # Affichage des résultats
            st.markdown('<div class="success-box">', unsafe_allow_html=True)
            st.success("✅ Analyse terminée avec succès!")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Affichage structuré des résultats
            st.markdown("## 📊 Résultats de l'Analyse")
            
            with st.expander("📋 Voir le rapport complet", expanded=True):
                st.markdown(analysis_result)
            
            # Téléchargement du rapport
            st.download_button(
                label="💾 Télécharger le rapport",
                data=analysis_result,
                file_name="analyse_vulnerabilites.md",
                mime="text/markdown"
            )

    except Exception as e:
        st.error(f"❌ Erreur lors du traitement: {str(e)}")
    finally:
        # Nettoyage du fichier temporaire
        os.unlink(tmp_path)

else:
    # Message d'accueil
    st.info("""
    👋 **Bienvenue dans l'analyseur de vulnérabilités!**
    
    Cette application vous permet d'automatiser l'analyse des vulnérabilités de sécurité détectées dans vos rapports d'audit.
    
    **Fonctionnalités:**
    - Extraction automatique des données DOCX
    - Analyse par IA spécialisée en cybersecurity
    - Génération de rapports structurés
    - Interface intuitive et ergonomique
    
    ➡️ **Commencez par télécharger votre rapport dans la zone ci-dessus!**
    """)
    
    # Section d'exemple
    with st.expander("📚 Exemple de format attendu"):
        st.code("""
        6.3. Vulnérabilités découvertes
        
        • [CRITIQUE] CVE-2024-1234: Vulnerability in XYZ component
          Description: Remote code execution vulnerability
          Solution: Apply patch version 2.1.4
          
        • [MOYEN] Weak password policy
          Description: Password complexity requirements not enforced
          Solution: Implement strong password policy
        """)

# Pied de page
st.markdown("---")
st.caption("🛡️ Outil d'analyse cybersecurity - Powered by Ollama AI")