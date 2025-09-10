import streamlit as st
import requests
from docx import Document
import re
import tempfile
import os
import json
from typing import Dict, List, Optional
import time

# Configuration de la page
st.set_page_config(
    page_title="Vérificateur de Cohérence des Vulnérabilités",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé
st.markdown("""
<style>
    .vulnerability-block {
        background-color: #F8F9FA;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #6c757d;
        margin: 15px 0;
    }
    .consistent {
        border-left: 5px solid #28a745 !important;
        background-color: #E8F5E9 !important;
    }
    .inconsistent {
        border-left: 5px solid #dc3545 !important;
        background-color: #FFEBEE !important;
    }
    .partial {
        border-left: 5px solid #ffc107 !important;
        background-color: #FFF3CD !important;
    }
    .section-box {
        background-color: white;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #E0E0E0;
        margin: 10px 0;
        font-size: 14px;
        max-height: 300px;
        overflow-y: auto;
    }
</style>
""", unsafe_allow_html=True)

def extract_vulnerabilities_properly(docx_path: str) -> List[Dict]:
    """
    Proper extraction that handles the complex table structure correctly.
    Each vulnerability has its own table structure.
    """
    doc = Document(docx_path)
    vulnerabilities = {}
    current_vuln_id = None
    
    # First, let's find all vulnerability titles in paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        # Look for vulnerability titles (headed with ##### or containing V_XXX)
        if text.startswith('##### ') or re.search(r'V_\d{2,3}', text):
            vuln_match = re.search(r'(V_\d{2,3})', text)
            if vuln_match:
                vuln_id = vuln_match.group(1)
                if vuln_id not in vulnerabilities:
                    vulnerabilities[vuln_id] = {
                        'id': vuln_id,
                        'title': text.replace('##### ', '').replace(vuln_id, '').strip(),
                        'constats': '',
                        'preuves': '',
                        'recommandation': '',
                        'impacts': '',
                        'niveau_criticite': '',
                        'elements_impactes': ''
                    }
    
    # Now process tables to extract content
    for table_idx, table in enumerate(doc.tables):
        current_vuln_id = None
        table_vuln_data = {}
        
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) >= 2:
                left_text = row.cells[0].text.strip()
                right_text = row.cells[1].text.strip()
                
                # Check if this row contains a vulnerability ID
                vuln_match = re.search(r'(V_\d{2,3})', left_text) or re.search(r'(V_\d{2,3})', right_text)
                if vuln_match:
                    current_vuln_id = vuln_match.group(1)
                    # Initialize if not exists
                    if current_vuln_id not in vulnerabilities:
                        vulnerabilities[current_vuln_id] = {
                            'id': current_vuln_id,
                            'title': right_text if vuln_match.group(1) in left_text else left_text,
                            'constats': '',
                            'preuves': '',
                            'recommandation': '',
                            'impacts': '',
                            'niveau_criticite': '',
                            'elements_impactes': ''
                        }
                    continue
                
                # If we have a current vulnerability, map the content
                if current_vuln_id:
                    section_map = {
                        'constats': ['constats', 'findings', 'observation'],
                        'preuves': ['preuves', 'evidence', 'proof'],
                        'recommandation': ['recommandation', 'recommendation', 'solution'],
                        'impacts': ['impacts', 'consequence', 'effect'],
                        'niveau_criticite': ['niveau de criticité', 'criticité', 'severity', 'risk level'],
                        'elements_impactes': ['éléments impactés', 'elements impactes', 'affected elements']
                    }
                    
                    for field, patterns in section_map.items():
                        if any(pattern in left_text.lower() for pattern in patterns):
                            # Only add if not already present (avoid duplicates)
                            if not vulnerabilities[current_vuln_id][field]:
                                vulnerabilities[current_vuln_id][field] = right_text
                            break
    
    # Convert to list, filter empty, and sort
    vuln_list = [v for v in vulnerabilities.values() if any(v[field] for field in ['constats', 'preuves', 'recommandation'])]
    vuln_list.sort(key=lambda x: x['id'])
    
    return vuln_list
def analyze_vulnerability_with_ai(vuln_data: Dict, model_choice: str) -> Dict:
    """
    Analyse la cohérence d'une vulnérabilité avec Ollama
    """
    prompt = f"""

ANALYSE DE COHÉRENCE D'UNE VULNÉRABILITÉ DE SÉCURITÉ

TU ES UN EXPERT EN CYBERSÉCURITÉ ET EN ANALYSE DE RAPPORTS D'AUDIT. 
TON RÔLE EST D'ANALYSER LA COHÉRENCE INTERNE D'UNE VULNÉRABILITÉ.

VULNÉRABILITÉ: {vuln_data.get('id', 'N/A')} - {vuln_data.get('title', 'Sans titre')}

INFORMATIONS DISPONIBLES:

CONSTATS (Description du problème):
{vuln_data.get('constats', 'Non spécifié')}

PREUVES (Éléments de preuve):
{vuln_data.get('preuves', 'Non spécifié')}

IMPACTS (Conséquences potentielles):
{vuln_data.get('impacts', 'Non spécifié')}

NIVEAU DE CRITICITÉ:
{vuln_data.get('niveau_criticite', 'Non spécifié')}

RECOMMANDATION (Solution proposée):
{vuln_data.get('recommandation', 'Non spécifié')}

TA MISSION:
1. ✅ Vérifier si les PREUVES supportent logiquement les CONSTATS
2. ✅ Vérifier si la RECOMMANDATION addressent correctement le problème des CONSTATS
3. ✅ Vérifier si le NIVEAU DE CRITICITÉ est approprié compte tenus des IMPACTS
4. 💯 Donner un score global de cohérence (0-100)
5. 🎯 Proposer des améliorations si nécessaire

RÉPONDS UNIQUEMENT en JSON avec ce format exact:
{{
  "verdict_global": "COHÉRENT" ou "PARTIELLEMENT COHÉRENT" ou "INCOHÉRENT",
  "score_cohérence": 0-100,
  "explication_globale": "Explication concise en 2-3 phrases",
  "points_forts": ["Point fort 1", "Point fort 2"],
  "points_amelioration": ["Point d'amélioration 1", "Point d'amélioration 2"],
  "details_analyse": {{
    "preuves_vs_constats": "Analyse détaillée de la cohérence preuves/constats",
    "recommandation_vs_constats": "Analyse détaillée de la pertinence des recommandations",
    "criticite_vs_impacts": "Analyse détaillée de l'adéquation criticité/impacts"
  }}
}}

Sois précis, technique et constructif dans ton analyse.
    """

    payload = {
        "model": model_choice,
        "prompt": prompt,
        "stream": False,   # on force le non-streaming
        "options": {
            "temperature": 0.1,
            "top_p": 0.9
        }
    }

    try:
        response = requests.post("http://localhost:11434/api/generate", json=payload, timeout=120)
        response.raise_for_status()

        # 🛠 Debug : voir ce que Ollama renvoie réellement
        raw_text = response.text.strip()
        print("RAW RESPONSE FROM OLLAMA:", raw_text[:500])  # affiche les 500 premiers caractères

        # ⚡ Certains modèles renvoient "```json ... ```", on nettoie
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]
        raw_text = raw_text.strip()

        # ✅ Essayer de parser le JSON
        try:
            return json.loads(raw_text)
        except json.JSONDecodeError:
            return {
                "error": "Réponse IA invalide (pas du JSON)",
                "raw_response": raw_text,
                "verdict_global": "ERREUR",
                "score_cohérence": 0,
                "explication_globale": "Impossible de parser la réponse IA"
            }

    except Exception as e:
        return {
            "error": f"Erreur lors de l'appel à Ollama: {str(e)}",
            "verdict_global": "ERREUR",
            "score_cohérence": 0,
            "explication_globale": "Impossible d'analyser avec l'IA"
        }

# Interface principale
st.markdown('<h1 style="text-align: center; color: #1E88E5;">🔍 Vérificateur de Cohérence des Vulnérabilités</h1>', unsafe_allow_html=True)
st.markdown("***Analyse IA des rapports d'audit de sécurité***")

# Sidebar
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/checklist.png", width=100)
    st.title("Configuration")
    
    model_choice = st.selectbox(
        "Modèle Ollama:",
        ["llama3.1:latest", "mistral:latest", "llama3:latest"],
        index=0,
        help="Choisissez le modèle pour l'analyse IA"
    )

# Zone de téléchargement
uploaded_file = st.file_uploader("📤 Téléchargez votre rapport d'audit (DOCX)", type="docx")

if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    try:
        # Extract vulnerabilities with new parser
        with st.spinner("🔍 Extraction des vulnérabilités en cours..."):
            vulnerabilities = extract_vulnerabilities_properly(tmp_path)
        
        if not vulnerabilities:
            st.error("Aucune vulnérabilité n'a pu être extraite.")
            st.stop()
        
        st.success(f"✅ {len(vulnerabilities)} vulnérabilité(s) trouvée(s)")
        
        # Select vulnerability to analyze
        vuln_options = [f"{vuln.get('id', 'N/A')}: {vuln.get('title', 'Sans titre')[:50]}..." for vuln in vulnerabilities]
        selected_index = st.selectbox("Sélectionnez une vulnérabilité à analyser:", range(len(vuln_options)), format_func=lambda x: vuln_options[x])
        
        selected_vuln = vulnerabilities[selected_index]
        
        # Display selected vulnerability
        st.markdown(f"## 🔍 {selected_vuln.get('id', 'N/A')} - {selected_vuln.get('title', 'Sans titre')}")
        
        # Display in two columns
        col1, col2 = st.columns(2)
        
        with col1:
            if selected_vuln.get('constats'):
                st.markdown("**📋 CONSTATS**")
                st.markdown(f'<div class="section-box">{selected_vuln["constats"]}</div>', unsafe_allow_html=True)
            
            if selected_vuln.get('preuves'):
                st.markdown("**🔎 PREUVES**")
                st.markdown(f'<div class="section-box">{selected_vuln["preuves"]}</div>', unsafe_allow_html=True)
        
        with col2:
            if selected_vuln.get('impacts'):
                st.markdown("**💥 IMPACTS**")
                st.markdown(f'<div class="section-box">{selected_vuln["impacts"]}</div>', unsafe_allow_html=True)
            
            if selected_vuln.get('niveau_criticite'):
                st.markdown("**⚠️ NIVEAU DE CRITICITÉ**")
                st.markdown(f'<div class="section-box">{selected_vuln["niveau_criticite"]}</div>', unsafe_allow_html=True)
            
            if selected_vuln.get('recommandation'):
                st.markdown("**✅ RECOMMANDATION**")
                st.markdown(f'<div class="section-box">{selected_vuln["recommandation"]}</div>', unsafe_allow_html=True)
        
        # Analysis button
        if st.button("🤖 Analyser avec l'IA", type="primary", use_container_width=True):
            # Test Ollama connection
            try:
                test_response = requests.get("http://localhost:11434/api/tags", timeout=10)
                if test_response.status_code != 200:
                    st.error("❌ Ollama n'est pas accessible")
                    st.stop()
                else:
                    st.success("✅ Ollama est connecté!")
            except:
                st.error("❌ Impossible de se connecter à Ollama")
                st.stop()
            
            # Perform AI analysis
            with st.spinner("🧠 Analyse IA en cours (2-3 minutes)..."):
                analysis_result = analyze_vulnerability_with_ai(selected_vuln, model_choice)
            
            # Display results
            if "error" in analysis_result:
                st.error(analysis_result["error"])
            else:
                st.markdown("## 📋 Résultats de l'Analyse IA")
                
                verdict = analysis_result.get('verdict_global', 'INCONNU')
                score = analysis_result.get('score_cohérence', 0)
                
                if verdict == "COHÉRENT":
                    st.success(f"**Verdict: {verdict}** - Score: {score}/100")
                elif verdict == "PARTIELLEMENT COHÉRENT":
                    st.warning(f"**Verdict: {verdict}** - Score: {score}/100")
                else:
                    st.error(f"**Verdict: {verdict}** - Score: {score}/100")
                
                st.info(f"**Explication:** {analysis_result.get('explication_globale', '')}")
                
                with st.expander("📊 Détails de l'analyse"):
                    st.json(analysis_result)

    except Exception as e:
        st.error(f"❌ Erreur lors du traitement: {str(e)}")
    finally:
        os.unlink(tmp_path)

else:
    st.info("""
    👋 **Bienvenue dans le Vérificateur de Cohérence IA!**
    
    **Fonctionnalités principales:**
    - 🔍 Extraction automatique des vulnérabilités depuis les rapports DOCX
    - 🤖 Analyse intelligente de la cohérence entre les sections
    - ✅ Validation des recommandations de sécurité
    - 📊 Scoring de qualité et suggestions d'amélioration
    
    **Comment utiliser:**
    1. 📤 Téléchargez un rapport d'audit au format DOCX
    2. 🔍 Sélectionnez une vulnérabilité à analyser
    3. 🤖 Lancez l'analyse IA
    4. 📊 Consultez les résultats détaillés
    """)